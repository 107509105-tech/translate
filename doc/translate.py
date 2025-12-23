"""
Bilingual Document Translation Tool
Translates Chinese Word documents to English while preserving the original Chinese text.
Uses Google Translator for translation with support for custom term dictionaries.
"""

import os
import re
import json
import time
from pathlib import Path
from copy import deepcopy
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.table import Table
from deep_translator import GoogleTranslator


# ==================== CONFIGURATION ====================

class TranslationConfig:
    """Central configuration for translation settings"""

    # Translator Settings (Google Translator)
    TRANSLATOR = GoogleTranslator(source='zh-TW', target='en')

    # File Paths
    PCB_TERMS_FILE = "pcb_terms_from_pdf.json"
    FIXED_TRANSLATION_FILE = Path('data/fixed_translation.json')
    
    # Font Settings
    DEFAULT_FONT = 'Times New Roman'
    DEFAULT_FONT_SIZE = 11
    HEADER_FOOTER_CHINESE_SIZE = 10
    HEADER_FOOTER_ENGLISH_SIZE = 6
    TABLE_HEADER_ENGLISH_SIZE = 8
    
    # Translation Settings
    FLOWCHART_TEXTBOX_THRESHOLD = 3
    TABLE_ENGLISH_FONT_RATIO = 0.82 


# ==================== STATE MANAGEMENT ====================

class DocumentState:
    """Manages global state during document translation"""
    
    def __init__(self):
        self.continuous_abnormal_groups = []
        self.current_group = None
        self.translated_group_ids = set()
        self.flowchart_original_tables = []
        self.pcb_terms_dict = {}
        self.fixed_translation_map = {}
    
    def load_translation_dictionaries(self):
        """Load PCB terms and fixed translation mappings"""
        # Load PCB terms
        # try:
        #     with open(TranslationConfig.PCB_TERMS_FILE, 'r', encoding='utf-8') as f:
        #         self.pcb_terms_dict = json.load(f)
        #     print(f"Loaded {len(self.pcb_terms_dict)} PCB terms")
        # except FileNotFoundError:
        #     print(f"Warning: {TranslationConfig.PCB_TERMS_FILE} not found, using Google Translator only")
        #     self.pcb_terms_dict = {}
        
        # Load fixed translations
        if TranslationConfig.FIXED_TRANSLATION_FILE.exists():
            with open(TranslationConfig.FIXED_TRANSLATION_FILE, 'r', encoding='utf-8') as f:
                self.fixed_translation_map = json.load(f)
            print(f"Loaded {len(self.fixed_translation_map)} fixed translations")
    
    def reset(self):
        """Reset state for new document"""
        self.continuous_abnormal_groups = []
        self.current_group = None
        self.translated_group_ids = set()
        self.flowchart_original_tables = []


# ==================== TEXT DETECTION UTILITIES ====================

class TextDetector:
    """Utilities for detecting text patterns"""
    
    @staticmethod
    def is_chinese(text):
        """Check if text contains Chinese characters"""
        return bool(re.search('[\u4e00-\u9fff]', text))
    
    @staticmethod
    def get_step_number(text):
        """Extract step number from text (e.g., '1.2.3', '1.2', '1.')"""
        if isinstance(text, str):
            raw = text
        else:
            raw = ''.join(run.text for run in text.runs)
        
        number_patterns = [
            r'^(\d+\.\d+\.\d+\.\d+)',  # 1.2.3.4
            r'^(\d+\.\d+.\d+)',         # 1.2.3
            r'^(\d+\.\d+)',             # 1.2
            r'^(\d+\.)',                # 1.
        ]
        
        for pattern in number_patterns:
            match = re.match(pattern, raw.lstrip())
            if match:
                return match.group(1).rstrip('.')
        
        return ""
    
    @staticmethod
    def check_colon_format(text):
        """
        Check if text has colon format
        Returns: (has_colon, has_content_after_colon, colon_part, content_part)
        """
        colon_match = re.match(r'^([^:：]+)[：:](.*)$', text.strip())
        if colon_match:
            before_colon = colon_match.group(1).strip()
            after_colon = colon_match.group(2).strip()
            return True, bool(after_colon), before_colon, after_colon
        return False, False, "", ""
    
    @staticmethod
    def has_long_spaces_in_runs(paragraph):
        """Detect leading spaces in paragraph"""
        raw_text = ''.join(run.text for run in paragraph.runs)
        
        if not raw_text:
            return False, 0
        
        leading_spaces = len(raw_text) - len(raw_text.lstrip(' \t'))
        
        if leading_spaces > 0:
            space_count = sum(4 if char == '\t' else 1 for char in raw_text[:leading_spaces])
            return True, space_count
        
        return False, 0


# ==================== WORD DOCUMENT UTILITIES ====================

class WordDocumentHelper:
    """Helper functions for Word document manipulation"""
    
    @staticmethod
    def has_picture(run):
        """Check if run contains image or drawing object"""
        if run._element.findall('.//' + qn('w:drawing')):
            return True
        if run._element.findall('.//' + qn('w:pict')):
            return True
        return False
    
    @staticmethod
    def clear_paragraph_text_keep_images(paragraph):
        """Clear text in paragraph while preserving images"""
        for run in paragraph.runs:
            if not WordDocumentHelper.has_picture(run):
                run.text = ""
    
    @staticmethod
    def is_special_format(run):
        """Check if run has special formatting (bold, color, images, etc.)"""
        font = run.font
        
        if font.bold or font.italic or font.underline:
            return True
        
        if font.color and font.color.rgb and str(font.color.rgb) not in ('000000', None, ''):
            return True
        
        if font.highlight_color:
            return True
        
        if WordDocumentHelper.has_picture(run):
            return True
        
        return False
    
    @staticmethod
    def add_english_below(paragraph, english_text, font_size=None, 
                          font_name=None, alignment=None):
        """Add English paragraph below Chinese paragraph for bilingual display"""
        if font_name is None:
            font_name = TranslationConfig.DEFAULT_FONT
        
        parent = paragraph._element.getparent()
        
        if parent is None:
            return None
        
        para_index = parent.index(paragraph._element)
        new_p = OxmlElement('w:p')
        parent.insert(para_index + 1, new_p)
        
        eng_para = Paragraph(new_p, paragraph._parent)
        
        # Set alignment
        if alignment == 'center':
            eng_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            eng_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == 'justify':
            eng_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif alignment == 'left':
            eng_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif paragraph.alignment is not None:
            eng_para.alignment = paragraph.alignment
        else:
            eng_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        run = eng_para.add_run(english_text)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        
        # Set font size
        if font_size is not None:
            run.font.size = Pt(font_size)
        elif paragraph.runs and paragraph.runs[0].font.size:
            run.font.size = paragraph.runs[0].font.size
        else:
            run.font.size = Pt(TranslationConfig.DEFAULT_FONT_SIZE)
        
        return eng_para
    
    @staticmethod
    def set_paragraph_font_size(paragraph, font_size_pt, target='all'):
        """Set font size for paragraph runs"""
        if not paragraph.runs:
            return
        
        half_points = font_size_pt * 2
        
        # Remove style binding for headers/footers
        pPr = paragraph._p.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                pPr.remove(pStyle)
        
        full_text = paragraph.text
        paragraph_has_chinese = TextDetector.is_chinese(full_text)
        
        should_adjust = (
            target == 'all' or
            (target == 'chinese' and paragraph_has_chinese) or
            (target == 'english' and not paragraph_has_chinese)
        )
        
        if not should_adjust:
            return
        
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            
            run.font.size = Pt(font_size_pt)
            
            rPr = run._r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                run._r.insert(0, rPr)
            
            for tag in ('w:sz', 'w:szCs'):
                elem = rPr.find(qn(tag))
                if elem is None:
                    elem = OxmlElement(tag)
                    rPr.insert(0, elem)
                elem.set(qn('w:val'), str(half_points))
    
    @staticmethod
    def find_ancestor_with_tag(element, target_tag):
        """Find ancestor element with specific tag"""
        parent = element.getparent()
        while parent is not None:
            if parent.tag == target_tag:
                return parent
            parent = parent.getparent()
        return None
    
    @staticmethod
    def insert_page_break_after(element):
        """Insert page break after specified element"""
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r.append(br)
        p.append(r)
        parent = element.getparent()
        insert_index = parent.index(element) + 1
        parent.insert(insert_index, p)
        return p


# ==================== TRANSLATION ENGINE ====================

class TranslationEngine:
    """Handles translation logic using Google Translator and dictionaries"""

    def __init__(self, state: DocumentState):
        self.state = state

    def translate_text(self, text):
        """Main translation function with dictionary lookup and Google Translator fallback"""
        if not text or not text.strip() or not TextDetector.is_chinese(text):
            return text
        
        text_stripped = text.strip()
        
        # Check fixed translation map first
        if text_stripped in self.state.fixed_translation_map:
            return self.state.fixed_translation_map[text_stripped]
        
        # Check PCB terms dictionary
        translation = self._check_pcb_terms(text_stripped)
        if translation:
            return translation

        # Use Google Translator for translation
        return self._translate_with_llm(text)
    
    def _check_pcb_terms(self, text):
        """Check PCB terms dictionary for translation"""
        for english_key, term_data in self.state.pcb_terms_dict.items():
            if not isinstance(term_data, dict):
                continue
            
            traditional = term_data.get('traditional', '')
            
            # Exact match
            if traditional == text:
                print(f'[JSON Exact Match] {text} → {english_key}')
                return english_key
            
            # Partial match (substring in comma-separated items)
            items = re.split(r'[，,]', traditional)
            for item in items:
                if item.strip() == text:
                    print(f'[JSON Partial Match] {text} (in {traditional}) → {english_key}')
                    return english_key
        
        return None
    
    def _translate_with_llm(self, text):
        """Translate text using Google Translator"""
        try:
            translated = TranslationConfig.TRANSLATOR.translate(text)
            return translated.strip()
        except Exception as e:
            print(f"Translation error for '{text}': {e}")
            return text


# ==================== PARAGRAPH PROCESSING ====================

class ParagraphProcessor:
    """Handles paragraph grouping and merging logic"""
    
    def __init__(self, state: DocumentState):
        self.state = state
    
    def record_long_space_paragraph(self, paragraph, para_index=None):
        """Record paragraphs with leading spaces for grouping"""
        has_abnormal, count = TextDetector.has_long_spaces_in_runs(paragraph)
        step_number = TextDetector.get_step_number(paragraph)

        # If paragraph has no leading spaces, end current group
        if not has_abnormal:
            if self.state.current_group is not None:
                self.state.current_group["merged_text"] = self.merge_group_text(
                    self.state.current_group["paragraphs"]
                )
                self.state.continuous_abnormal_groups.append(self.state.current_group)
                self.state.current_group = None
            return

        if has_abnormal and para_index is not None:
            if step_number:
                # Has step number → end old group, start new group
                if self.state.current_group is not None:
                    self.state.current_group["merged_text"] = self.merge_group_text(
                        self.state.current_group["paragraphs"]
                    )
                    self.state.continuous_abnormal_groups.append(self.state.current_group)
                    self.state.current_group = None

                # Start new group
                self.state.current_group = {
                    "group_id": len(self.state.continuous_abnormal_groups) + 1,
                    "paragraphs": []
                }
                self.state.current_group["paragraphs"].append({
                    "para_index": para_index,
                    'para': paragraph,
                    "full_text": paragraph.text.strip(),
                    "space_count": count
                })
            else:
                # No step number but has spaces → add to current group or create new group
                # These are either continuation lines after a numbered paragraph,
                # or consecutive indented paragraphs that should be merged
                if self.state.current_group is not None:
                    self.state.current_group["paragraphs"].append({
                        "para_index": para_index,
                        'para': paragraph,
                        "full_text": paragraph.text.strip(),
                        "space_count": count
                    })
                else:
                    # No current group - create one for consecutive indented paragraphs
                    self.state.current_group = {
                        "group_id": len(self.state.continuous_abnormal_groups) + 1,
                        "paragraphs": []
                    }
                    self.state.current_group["paragraphs"].append({
                        "para_index": para_index,
                        'para': paragraph,
                        "full_text": paragraph.text.strip(),
                        "space_count": count
                    })
    
    def merge_group_text(self, paragraphs_list):
        """Merge grouped paragraph text"""
        if not paragraphs_list:
            return ""
        
        lines = [p["full_text"].lstrip() for p in paragraphs_list]
        
        is_step_group = any(TextDetector.get_step_number(p["full_text"]) 
                           for p in paragraphs_list)
        
        if is_step_group:
            result = lines[0]
            for line in lines[1:]:
                result += line.lstrip()
            return result
        else:
            return " ".join(lines)
    
    def finalize_groups(self):
        """Finalize the current group"""
        if self.state.current_group is not None:
            self.state.current_group["merged_text"] = self.merge_group_text(
                self.state.current_group["paragraphs"]
            )
            self.state.continuous_abnormal_groups.append(self.state.current_group)


# ==================== BILINGUAL TRANSLATOR ====================

class BilingualTranslator:
    """Main translation coordinator for bilingual output"""
    
    def __init__(self, state: DocumentState):
        self.state = state
        self.engine = TranslationEngine(state)
        self.processor = ParagraphProcessor(state)
    
    def translate_paragraph(self, paragraph, para_index=None):
        """Translate paragraph in bilingual mode (keep Chinese, add English below)"""
        if not paragraph.text.strip() or not TextDetector.is_chinese(paragraph.text):
            return

        # Check if paragraph belongs to a grouped section
        belonging_group = None
        is_first_para_in_group = False

        for group in self.state.continuous_abnormal_groups:
            if group["group_id"] in self.state.translated_group_ids:
                continue
            for item in group["paragraphs"]:
                if item["para_index"] == para_index:
                    belonging_group = group
                    is_first_para_in_group = (
                        item["para_index"] == group["paragraphs"][0]["para_index"]
                    )
                    break
            if belonging_group:
                break

        # Skip non-first paragraphs in groups
        if belonging_group and not is_first_para_in_group:
            return

        # Handle grouped paragraphs
        if belonging_group and is_first_para_in_group:
            self._translate_grouped_paragraphs(paragraph, belonging_group)
            return

        # Handle regular paragraphs
        self._translate_regular_paragraph(paragraph)
    
    def _translate_grouped_paragraphs(self, paragraph, group):
        """Translate grouped paragraphs as one unit"""
        merged_chinese = self.processor.merge_group_text(group["paragraphs"])
        first_full_text = group["paragraphs"][0]["full_text"]
        step_number = TextDetector.get_step_number(first_full_text)
        
        # Preserve indentation
        raw_text = paragraph.text
        leading_spaces = len(raw_text) - len(raw_text.lstrip())
        indent = " " * leading_spaces
        
        # Update first paragraph with merged content
        has_images = any(WordDocumentHelper.has_picture(run) for run in paragraph.runs)
        
        if has_images:
            text_runs = [run for run in paragraph.runs 
                        if not WordDocumentHelper.has_picture(run)]
            for run in text_runs:
                run.text = ""
            paragraph.add_run(merged_chinese)
        else:
            for run in paragraph.runs:
                run.text = ""
            paragraph.add_run(indent + merged_chinese)
        
        # Remove step number for translation
        if step_number:
            content_start = merged_chinese.find(step_number) + len(step_number)
            while content_start < len(merged_chinese) and \
                  merged_chinese[content_start] in '. :：\u3000\t ':
                content_start += 1
            pure_content = merged_chinese[content_start:]
        else:
            pure_content = merged_chinese
        
        # Translate
        translated_full = self.engine.translate_text(pure_content)
        final_english = indent + translated_full.strip()
        
        # Add English paragraph
        WordDocumentHelper.add_english_below(paragraph, final_english)
        
        # Clear subsequent paragraphs
        for item in group["paragraphs"][1:]:
            WordDocumentHelper.clear_paragraph_text_keep_images(item["para"])
        
        self.state.translated_group_ids.add(group["group_id"])
    
    def _translate_regular_paragraph(self, paragraph):
        """Translate regular (non-grouped) paragraph"""
        raw_text = paragraph.text
        stripped_text = raw_text.strip()
        
        # Check colon format
        has_colon, has_content, colon_part, _ = TextDetector.check_colon_format(stripped_text)
        
        if has_colon and not has_content:
            # Format 1: Colon with no content → add translation in brackets
            self._translate_colon_no_content(paragraph, raw_text, colon_part)
        elif has_colon and has_content:
            # Format 2: Colon with content → add English below
            self._translate_colon_with_content(paragraph, raw_text, stripped_text)
        else:
            # Format 3: No colon → add English below
            self._translate_no_colon(paragraph, raw_text, stripped_text)
    
    def _translate_colon_no_content(self, paragraph, raw_text, colon_part):
        """Handle colon format with no content after colon"""
        translated = self.engine.translate_text(colon_part)
        has_images = any(WordDocumentHelper.has_picture(run) for run in paragraph.runs)
        
        original_colon = '：' if '：' in raw_text else ':'
        leading_spaces = len(raw_text) - len(raw_text.lstrip())
        indent = " " * leading_spaces
        new_text = f"{indent}{colon_part}({translated}){original_colon}"
        
        if has_images:
            for run in paragraph.runs:
                if not WordDocumentHelper.has_picture(run):
                    run.text = ""
            paragraph.add_run(new_text)
        else:
            for run in paragraph.runs:
                run.text = ""
            paragraph.add_run(new_text)
    
    def _translate_colon_with_content(self, paragraph, raw_text, stripped_text):
        """Handle colon format with content after colon"""
        step_number = TextDetector.get_step_number(stripped_text)
        leading_spaces = len(raw_text) - len(raw_text.lstrip())
        indent = " " * leading_spaces
        
        if step_number:
            pure_content = self._extract_content_after_number(stripped_text, step_number)
            translated = self.engine.translate_text(pure_content)
            final_english = f"{indent}{translated}"
        else:
            translated = self.engine.translate_text(stripped_text)
            final_english = f"{indent}{translated}"
        
        WordDocumentHelper.add_english_below(paragraph, final_english)
    
    def _translate_no_colon(self, paragraph, raw_text, stripped_text):
        """Handle paragraph with no colon"""
        step_number = TextDetector.get_step_number(stripped_text)
        leading_spaces = len(raw_text) - len(raw_text.lstrip())
        indent = " " * leading_spaces
        
        if step_number:
            pure_content = self._extract_content_after_number(stripped_text, step_number)
            translated = self.engine.translate_text(pure_content)
            final_english = f"{indent}{translated}"
        else:
            translated = self.engine.translate_text(stripped_text)
            final_english = f"{indent}{translated}"
        
        WordDocumentHelper.add_english_below(paragraph, final_english)
    
    def _extract_content_after_number(self, text, step_number):
        """Extract content after step number"""
        content_start = text.find(step_number) + len(step_number)
        while content_start < len(text) and text[content_start] in '. :：\u3000\t ':
            content_start += 1
        return text[content_start:]
    
    def translate_table(self, table):
        """Translate table content in bilingual mode"""
        for row in table.rows:
            for cell in row.cells:
                paragraphs = [p for p in cell.paragraphs 
                            if p.text.strip() and TextDetector.is_chinese(p.text)]
                
                if not paragraphs:
                    continue
                
                has_numbers = any(TextDetector.get_step_number(p.text) for p in paragraphs)
                
                if has_numbers and len(paragraphs) > 1:
                    self._translate_numbered_table_cell(paragraphs)
                else:
                    for p in paragraphs:
                        self.translate_paragraph(p)
    
    def _translate_numbered_table_cell(self, paragraphs):
        """Translate table cell with multiple numbered items"""
        all_chinese = []
        all_numbers = []
        
        for p in paragraphs:
            stripped = p.text.strip()
            step_num = TextDetector.get_step_number(stripped)
            
            if step_num:
                pure_content = self._extract_content_after_number(stripped, step_num)
                all_chinese.append(pure_content)
                all_numbers.append(step_num)
            else:
                all_chinese.append(stripped)
                all_numbers.append("")
        
        merged_chinese = " ".join(all_chinese)
        translated_full = self.engine.translate_text(merged_chinese)
        
        translated_parts = [p.strip() for p in translated_full.split('.') if p.strip()]
        
        english_lines = []
        for i, num in enumerate(all_numbers):
            if i < len(translated_parts):
                if num:
                    english_lines.append(f"{num}.{translated_parts[i]}")
                else:
                    english_lines.append(translated_parts[i])
        
        if len(translated_parts) > len(all_numbers):
            english_lines.extend(translated_parts[len(all_numbers):])
        
        last_para = paragraphs[-1]
        combined_english = "\n".join(english_lines)
        WordDocumentHelper.add_english_below(last_para, combined_english)
    
    def translate_header_footer(self, doc):
        """Translate headers and footers"""
        for section in doc.sections:
            headers = [
                section.header,
                section.first_page_header,
                section.even_page_header
            ]
            footers = [
                section.footer,
                section.first_page_footer,
                section.even_page_footer
            ]
            
            for hf in headers + footers:
                if not hf or hf.is_linked_to_previous:
                    continue
                
                # Translate paragraphs
                for para in hf.paragraphs:
                    if para.text.strip() and TextDetector.is_chinese(para.text):
                        WordDocumentHelper.set_paragraph_font_size(
                            para, TranslationConfig.HEADER_FOOTER_CHINESE_SIZE, 'chinese'
                        )
                        print(f"Header/Footer: {para.text}")
                        translated = self.engine.translate_text(para.text)
                        if translated:
                            WordDocumentHelper.add_english_below(
                                para, translated, 
                                font_size=TranslationConfig.HEADER_FOOTER_ENGLISH_SIZE,
                                alignment='center'
                            )
                
                # Translate tables in header/footer
                for table in hf.tables:
                    processed_texts = set()
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                full_text = para.text.strip()
                                if full_text and TextDetector.is_chinese(full_text):
                                    WordDocumentHelper.set_paragraph_font_size(
                                        para, TranslationConfig.HEADER_FOOTER_CHINESE_SIZE, 'chinese'
                                    )
                                    if full_text in processed_texts:
                                        continue
                                    
                                    print(f"Header/Footer Table: {para.text}")
                                    translated = self.engine.translate_text(para.text)
                                    if translated:
                                        WordDocumentHelper.add_english_below(
                                            para, translated,
                                            font_size=TranslationConfig.TABLE_HEADER_ENGLISH_SIZE,
                                            alignment='center'
                                        )
                                        processed_texts.add(full_text)


# ==================== FLOWCHART HANDLER ====================

class FlowchartHandler:
    """Handles flowchart detection and translation"""
    
    def __init__(self, state: DocumentState):
        self.state = state
        self.engine = TranslationEngine(state)
    
    def count_textboxes_in_table(self, table):
        """Count textboxes in table"""
        count = 0
        for row in table.rows:
            for cell in row.cells:
                count += len(cell._element.findall('.//' + qn('w:txbxContent')))
        return count
    
    def is_flowchart_table(self, table):
        """Determine if table is a flowchart based on textbox count"""
        return self.count_textboxes_in_table(table) >= TranslationConfig.FLOWCHART_TEXTBOX_THRESHOLD
    
    def clone_and_translate_flowchart(self, table):
        """Clone flowchart table and translate the copy"""
        original_elem = table._element
        parent = original_elem.getparent()
        
        # Track original table
        self.state.flowchart_original_tables.append(original_elem)
        
        # Insert page break
        page_break_para = WordDocumentHelper.insert_page_break_after(original_elem)
        
        # Clone table
        cloned_elem = deepcopy(original_elem)
        insert_index = parent.index(page_break_para) + 1
        parent.insert(insert_index, cloned_elem)
        
        translated_table = Table(cloned_elem, table._parent)
        
        # Translate cloned table
        for row in translated_table.rows:
            for cell in row.cells:
                # Translate paragraphs
                for para in cell.paragraphs:
                    if para.text.strip() and TextDetector.is_chinese(para.text):
                        translated = self.engine.translate_text(para.text.strip())
                        for run in para.runs:
                            run.text = ""
                        para.add_run(translated)
                
                # Translate textboxes
                # textboxes = cell._element.findall('.//' + qn('w:txbxContent'))
                # for textbox in textboxes:
                #     text_elements = textbox.findall('.//' + qn('w:t'))
                #     for text_elem in text_elements:
                #         if text_elem.text and text_elem.text.strip():
                #             translated_text = self.engine.translate_text(text_elem.text)
                #             if translated_text:
                #                 text_elem.text = translated_text

                textboxes = cell._element.findall('.//' + qn('w:txbxContent'))



                for textbox in textboxes:
                    # 先翻譯所有文字
                    text_elements = textbox.findall('.//' + qn('w:t'))
                    for text_elem in text_elements:
                        if text_elem.text and text_elem.text.strip():
                            original_text = text_elem.text
                            #print(original_text)
                            translated_text = self.engine.translate_text(original_text)
                            if translated_text:
                                text_elem.text = translated_text
                            
                            if TextDetector.is_chinese(original_text):
                                half_points = 11
                                # 調整統一套字體大小
                                for r in textbox.findall('.//' + qn('w:r')):
                                    rPr = r.find(qn('w:rPr'))
                                    if rPr is None:
                                        rPr = OxmlElement('w:rPr')
                                        r.append(rPr)
                                    sz = rPr.find(qn('w:sz'))
                                    if sz is None:
                                        sz = OxmlElement('w:sz')
                                        rPr.append(sz)
                                    sz.set(qn('w:val'), str(half_points))

                                # 調整行高                
                                for p in textbox.findall('.//' + qn('w:p')):
                                    pPr = p.find(qn('w:pPr'))
                                    if pPr is None:
                                        pPr = OxmlElement('w:pPr')
                                        p.insert(0, pPr)
                                    spacing = pPr.find(qn('w:spacing'))
                                    if spacing is None:
                                        spacing = OxmlElement('w:spacing')
                                        pPr.append(spacing)
                                    spacing.set(qn('w:line'), '130')
                                    spacing.set(qn('w:lineRule'), 'exact')
                                
                                # 調整對齊
                                for p in textbox.findall('.//' + qn('w:p')):
                                    pPr = p.find(qn('w:pPr'))
                                    if pPr is None:
                                        pPr = OxmlElement('w:pPr')
                                        p.insert(0, pPr)  # 必須 insert(0)

                                    # 確保有 w:jc，並設定為置中
                                    jc = pPr.find(qn('w:jc'))
                                    if jc is None:
                                        jc = OxmlElement('w:jc')
                                        pPr.append(jc)
                                    jc.set(qn('w:val'), 'center')
                    
                    return translated_table
                
    def translate_textboxes_in_doc(self, doc):
        """Translate all textboxes in document"""
        body = doc.element.body
        if body is None:
            return
        
        textboxes = body.findall('.//' + qn('w:txbxContent'))
        
        if not textboxes:
            return
        
        for textbox in textboxes:
            # Skip original flowchart tables
            ancestor_table = WordDocumentHelper.find_ancestor_with_tag(textbox, qn('w:tbl'))
            if ancestor_table in self.state.flowchart_original_tables:
                continue
            
            # Translate text
            text_elements = textbox.findall('.//' + qn('w:t'))
            for text_elem in text_elements:
                if text_elem.text and text_elem.text.strip():
                    original_text = text_elem.text
                    translated_text = self.engine.translate_text(original_text)
                    if translated_text:
                        text_elem.text = translated_text
                    
                    if TextDetector.is_chinese(original_text):
                        self._adjust_textbox_formatting(textbox)
    
    def _adjust_textbox_formatting(self, textbox):
        """Adjust font size and alignment for textbox"""
        half_points = 11
        
        # Adjust font size
        for r in textbox.findall('.//' + qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.append(rPr)
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = OxmlElement('w:sz')
                rPr.append(sz)
            sz.set(qn('w:val'), str(half_points))
        
        # Adjust line height
        for p in textbox.findall('.//' + qn('w:p')):
            pPr = p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p.insert(0, pPr)
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            spacing.set(qn('w:line'), '130')
            spacing.set(qn('w:lineRule'), 'exact')
        
        # Adjust alignment
        for p in textbox.findall('.//' + qn('w:p')):
            pPr = p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p.insert(0, pPr)
            
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'center')


# ==================== DOCUMENT FORMATTER ====================

class DocumentFormatter:
    """Handles document-wide formatting operations"""
    
    @staticmethod
    def shrink_table_english_font(table, ratio=None):
        """Shrink English text in tables"""
        if ratio is None:
            ratio = TranslationConfig.TABLE_ENGLISH_FONT_RATIO
        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if not text or TextDetector.is_chinese(text):
                        continue
                    
                    for run in paragraph.runs:
                        if not run.text.strip():
                            continue
                        
                        r = run._r
                        rPr = r.find(qn('w:rPr'))
                        if rPr is None:
                            rPr = OxmlElement('w:rPr')
                            r.insert(0, rPr)
                        
                        for tag in ('w:sz', 'w:szCs', 'w:szFarEast'):
                            elem = rPr.find(qn(tag))
                            if elem is None:
                                elem = OxmlElement(tag)
                                rPr.insert(0, elem)
                            
                            current_val = elem.get(qn('w:val'))
                            if current_val:
                                new_val = max(20, int(int(current_val) * ratio))
                            else:
                                new_val = max(20, int(22 * ratio))
                            
                            elem.set(qn('w:val'), str(new_val))
    
    @staticmethod
    def force_times_new_roman(doc):
        """Apply Times New Roman font to all English text"""
        def set_font(run):
            text = run.text.strip()
            if not text or TextDetector.is_chinese(text):
                return
            
            r = run._r
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            
            font_name = TranslationConfig.DEFAULT_FONT
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:cs'), font_name)
        
        # Apply to paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                set_font(run)
        
        # Apply to tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_font(run)
        
        # Apply to headers/footers
        for section in doc.sections:
            for hf in (section.header, section.footer,
                      section.first_page_header, section.first_page_footer,
                      section.even_page_header, section.even_page_footer):
                if not hf:
                    continue
                
                for paragraph in hf.paragraphs:
                    for run in paragraph.runs:
                        set_font(run)
                
                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    set_font(run)
        
        # Apply to textboxes
        for textbox in doc.element.body.findall('.//' + qn('w:txbxContent')):
            for run in textbox.findall('.//' + qn('w:r')):
                text_elements = run.findall('.//' + qn('w:t'))
                if not any(TextDetector.is_chinese(t.text or '') for t in text_elements):
                    rPr = run.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        run.insert(0, rPr)
                    
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.insert(0, rFonts)
                    
                    font_name = TranslationConfig.DEFAULT_FONT
                    rFonts.set(qn('w:ascii'), font_name)
                    rFonts.set(qn('w:hAnsi'), font_name)
                    rFonts.set(qn('w:cs'), font_name)
    
    @staticmethod
    def remove_empty_paragraphs(doc):
        """Remove empty paragraphs while preserving images"""
        paragraphs_to_remove = []
        
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                has_images = any(WordDocumentHelper.has_picture(run) 
                               for run in paragraph.runs)
                if not has_images:
                    paragraphs_to_remove.append(paragraph)
        
        for para in paragraphs_to_remove:
            p_element = para._element
            parent = p_element.getparent()
            if parent is not None:
                parent.remove(p_element)


# ==================== DOCUMENT TRANSLATOR ====================

class DocumentTranslator:
    """Main orchestrator for document translation"""
    
    def __init__(self):
        self.state = DocumentState()
        self.translator = BilingualTranslator(self.state)
        self.flowchart_handler = FlowchartHandler(self.state)
        self.formatter = DocumentFormatter()
    
    def translate_document(self, input_file, output_file):
        """Main translation workflow"""
        os.system('cls' if os.name == 'nt' else 'clear')
        
        # Reset state
        self.state.reset()
        self.state.load_translation_dictionaries()
        
        print("=" * 90)
        print(f"Loading file: {input_file}\n")
        doc = Document(input_file)

        # Step 1: Group paragraphs with leading spaces
        print("Analyzing paragraph structure...")
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip() and TextDetector.is_chinese(paragraph.text):
                self.translator.processor.record_long_space_paragraph(paragraph, para_index=i)
            
        self.translator.processor.finalize_groups()

        # Step 2: Translate paragraphs
        print("Translating paragraphs (bilingual mode)...")
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                self.translator.translate_paragraph(paragraph, para_index=i)
        
        # Step 3: Translate tables
        print("\nTranslating tables (bilingual mode)...")
        for table in doc.tables:
            if self.flowchart_handler.is_flowchart_table(table):
                print(f"Flowchart detected ({self.flowchart_handler.count_textboxes_in_table(table)} textboxes), cloning and translating...")
                self.flowchart_handler.clone_and_translate_flowchart(table)
            else:
                self.translator.translate_table(table)
        
        # Step 4: Translate headers/footers
        print("\nTranslating headers and footers...")
        self.translator.translate_header_footer(doc)
        
        # Step 5: Translate textboxes/flowcharts
        print("\nTranslating flowcharts and textboxes...")
        self.flowchart_handler.translate_textboxes_in_doc(doc)
        
        # Step 6: Format tables
        print("\nShrinking English font in tables...")
        for table in doc.tables:
            self.formatter.shrink_table_english_font(table)
        
        for section in doc.sections:
            for hf in (section.header, section.footer,
                      section.first_page_header, section.first_page_footer,
                      section.even_page_header, section.even_page_footer):
                if hf and hf.tables:
                    for t in hf.tables:
                        self.formatter.shrink_table_english_font(t)
        
        # Step 7: Apply font formatting
        print("Applying Times New Roman font...")
        self.formatter.force_times_new_roman(doc)
        
        # Step 8: Clean up
        print("Cleaning empty paragraphs...")
        self.formatter.remove_empty_paragraphs(doc)
        
        # Save
        print("=" * 90)
        print(f"Saving translated document → {output_file}")
        doc.save(output_file)
        print("Translation complete!")



# ==================== MAIN EXECUTION ====================

def main():
    """Main entry point"""
    print("Multilingual Document Translation Project (Bilingual Mode)\n")
    
    # Configure input/output
    # test_file = Path(r"I:\Multilingual Document Translation\W-QA-B002_OQC檢驗作業標準.docx")
    test_file = Path("document_cn.docx")
    # test_file = Path(r"Q:\14_人工智慧處\01_AI部\01_智慧代理應用課\01_個人資料夾\Ladem\for 竑軒\W-QA-A070 ver02 原子吸收光譜儀(280FS AA)操作指導書.docx")
    
    output_file = Path("out_bilingual_refactored2.docx")
    
    print(f"Input file: {test_file}")
    print(f"Output file: {output_file}\n")
    
    start_time = time.time()
    
    # Translate
    translator = DocumentTranslator()
    translator.translate_document(test_file, output_file)
    
    # Log results
    total_time = time.time() - start_time
    print(f"\nTotal time: {total_time:.2f} seconds ({total_time/60:.2f} minutes)")
    
    with open("翻譯時間紀錄.log", "a", encoding="utf-8") as f:
        f.write(f"{datetime.now():%Y-%m-%d %H:%M:%S} | [Refactored-Bilingual] | "
                f"Time: {total_time/60:.2f} minutes\n")


if __name__ == "__main__":
    main()