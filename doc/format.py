from docx import Document
import re

def find_merge_ranges(paragraphs):
    """
    規則：
    1. 以編號（如 4.9, 5.1）開頭的段落 = 區塊起點
    2. 以句末標點結尾的段落 = 區塊終點
    3. 起點到終點之間的所有段落合併成一段
    """
    number_pattern = re.compile(r'^\s*\d+\.\d+')  # 匹配 4.9, 5.1 等
    # 真正的句末標點（不包含括號類，因為內容可能繼續）
    end_marks = ('。', '！', '？')
    # 或者以 ) 結尾且前面是「圖X」的情況
    figure_pattern = re.compile(r'\(圖.+\)$|\（圖.+\）$')
    
    ranges = []
    start_idx = None
    
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # 遇到新編號
        if number_pattern.match(text):
            # 如果前一個區塊還沒結束，強制結束它
            if start_idx is not None and i - 1 > start_idx:
                ranges.append((start_idx, i - 1))
            start_idx = i
        
        # 檢查是否句子結束（句末標點 或 以「(圖X)」結尾）
        if start_idx is not None:
            is_end = text.endswith(end_marks) or figure_pattern.search(text)
            if is_end and i > start_idx:  # 至少跨兩段才需要合併
                ranges.append((start_idx, i))
                start_idx = None
    
    return ranges


def do_merge(doc, ranges):
    """執行合併（從後往前處理避免索引錯亂）"""
    ranges = sorted(ranges, key=lambda x: x[0], reverse=True)
    
    for start_idx, end_idx in ranges:
        # 合併文字，去掉每行首尾空白
        merged = ''.join(
            doc.paragraphs[i].text.strip()
            for i in range(start_idx, end_idx + 1)
        )
        
        # 更新第一段
        doc.paragraphs[start_idx].clear()
        doc.paragraphs[start_idx].add_run(merged)
        
        # 刪除後續段落
        for i in range(end_idx, start_idx, -1):
            p = doc.paragraphs[i]._element
            p.getparent().remove(p)
        
        print(f"合併段落 {start_idx}-{end_idx}: {merged[:50]}...")


# ========== 執行 ==========
doc = Document('/home/claude/worksheet.docx')

# 先預覽所有段落
print("=== 文檔段落預覽 ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f"[{i:2d}] {text[:60]}{'...' if len(text) > 60 else ''}")

print("\n=== 找到需要合併的範圍 ===")
ranges = find_merge_ranges(doc.paragraphs)
for start, end in ranges:
    print(f"  段落 {start} ~ {end}")

print("\n=== 執行合併 ===")
do_merge(doc, ranges)

output_path = '/mnt/user-data/outputs/worksheet_auto_merged.docx'
doc.save(output_path)
print(f"\n完成！儲存至: {output_path}")








from docx import Document
import re

doc = Document('your_file.docx')

# ========== 規則 1：句子未結束 ==========
# 如果段落不以句末標點結尾，就跟下一段合併
def merge_by_incomplete_sentence(paragraphs):
    """段落結尾不是。！？) 則視為未完成"""
    end_marks = ('。', '！', '？', ')', '）', '"', '」')
    
    i = 0
    while i < len(paragraphs) - 1:
        text = paragraphs[i].text.strip()
        if text and not text.endswith(end_marks):
            # 合併到下一段
            yield (i, i + 1)
        i += 1


# ========== 規則 2：編號段落 ==========
# 以編號開頭的段落，直到下一個編號前都屬於同一段
def merge_by_numbering(paragraphs):
    """如 4.9 開頭的段落，後續無編號的都合併"""
    number_pattern = re.compile(r'^\s*\d+\.\d+')
    
    start_idx = None
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if number_pattern.match(text):
            if start_idx is not None:
                yield (start_idx, i - 1)
            start_idx = i
    
    # 最後一個編號段落到文檔結尾
    if start_idx is not None:
        yield (start_idx, len(paragraphs) - 1)


# ========== 規則 3：連續縮排 ==========
# 開頭有空格縮排的連續段落合併
def merge_by_indentation(paragraphs):
    """連續以空格開頭的段落合併在一起"""
    start_idx = None
    
    for i, para in enumerate(paragraphs):
        text = para.text
        is_indented = text.startswith(' ') or text.startswith('\t')
        
        if is_indented and start_idx is None:
            start_idx = i
        elif not is_indented and start_idx is not None:
            if i - 1 > start_idx:  # 至少兩段才需要合併
                yield (start_idx, i - 1)
            start_idx = None


# ========== 規則 4：短行合併 ==========
# 連續的短段落（可能是被錯誤斷行）合併
def merge_short_lines(paragraphs, max_length=50):
    """連續短於 max_length 字元的段落合併"""
    start_idx = None
    
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        is_short = 0 < len(text) < max_length
        
        if is_short and start_idx is None:
            start_idx = i
        elif not is_short and start_idx is not None:
            if i - 1 > start_idx:
                yield (start_idx, i - 1)
            start_idx = None


# ========== 通用合併函數 ==========
def do_merge(doc, ranges):
    """執行合併，ranges 是 (start, end) 的列表，需從後往前處理"""
    ranges = sorted(ranges, key=lambda x: x[0], reverse=True)
    
    for start_idx, end_idx in ranges:
        if start_idx >= end_idx:
            continue
            
        # 合併文字
        merged = ''.join(
            doc.paragraphs[i].text.strip() 
            for i in range(start_idx, end_idx + 1)
        )
        
        # 更新第一段
        doc.paragraphs[start_idx].clear()
        doc.paragraphs[start_idx].add_run(merged)
        
        # 刪除後續段落
        for i in range(end_idx, start_idx, -1):
            p = doc.paragraphs[i]._element
            p.getparent().remove(p)


# ========== 使用範例 ==========
doc = Document('input.docx')

# 選擇你需要的規則
ranges = list(merge_by_incomplete_sentence(doc.paragraphs))
# 或: ranges = list(merge_by_numbering(doc.paragraphs))

do_merge(doc, ranges)
doc.save('output.docx')
