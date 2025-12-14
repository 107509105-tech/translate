from docx import Document

def merge_and_translate_paragraphs(doc_path, output_path, translate_func):
    doc = Document(doc_path)
    paragraphs = doc.paragraphs
    
    i = 0
    while i < len(paragraphs):
        # 檢測是否是被分割的短段落序列
        merged_text, merge_count = detect_and_merge(paragraphs, i)
        
        if merge_count > 1:
            # 翻譯合併後的文字
            translated = translate_func(merged_text)
            
            # 把結果寫入第一個段落，清空其他段落
            paragraphs[i].clear()
            paragraphs[i].add_run(translated)
            
            for j in range(1, merge_count):
                paragraphs[i + j].clear()
            
            i += merge_count
        else:
            # 正常翻譯單個段落
            full_text = ''.join(run.text for run in paragraphs[i].runs)
            if full_text.strip():
                translated = translate_func(full_text)
                paragraphs[i].clear()
                paragraphs[i].add_run(translated)
            i += 1
    
    doc.save(output_path)


def detect_and_merge(paragraphs, start_idx):
    """
    檢測從 start_idx 開始的連續短段落，判斷是否需要合併
    返回: (合併後的文字, 合併的段落數量)
    """
    texts = []
    count = 0
    
    for i in range(start_idx, len(paragraphs)):
        text = ''.join(run.text for run in paragraphs[i].runs).strip()
        
        if not text:
            break
            
        texts.append(text)
        count += 1
        
        # 判斷是否應該繼續合併
        combined = ''.join(texts)
        
        # 如果已經形成完整的「圖X」模式，停止合併
        if is_complete_figure_reference(combined):
            break
        
        # 如果當前段落已經夠長，不需要合併
        if len(text) > 5 and count == 1:
            break
        
        # 如果下一個段落不像是被分割的片段，停止
        if i + 1 < len(paragraphs):
            next_text = ''.join(run.text for run in paragraphs[i + 1].runs).strip()
            if not should_merge_next(combined, next_text):
                break
    
    if count > 1:
        return ''.join(texts), count
    else:
        return texts[0] if texts else '', 1


def is_complete_figure_reference(text):
    """判斷是否是完整的圖表引用"""
    import re
    # 匹配「圖一」到「圖九十九」等
    pattern = r'^圖[一二三四五六七八九十零百]+$'
    return bool(re.match(pattern, text))


def should_merge_next(current, next_text):
    """判斷是否應該合併下一個段落"""
    if not next_text or len(next_text) > 10:
        return False
    
    # 如果當前是「圖」，下一個是數字，應該合併
    if current == '圖' and next_text in '一二三四五六七八九十零百':
        return True
    
    # 如果當前是「圖X」，下一個還是數字，繼續合併
    if current.startswith('圖') and next_text in '一二三四五六七八九十零百':
        return True
    
    # 其他短片段也考慮合併
    if len(current) <= 3 and len(next_text) <= 3:
        return True
    
    return False
