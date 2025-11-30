#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
將中文 Word 文件翻譯成英文，並保留原始格式
"""

from docx import Document
from deep_translator import GoogleTranslator
from docx.oxml.ns import qn
import time

def translate_text(text, translator):
    """
    翻譯文本，處理空字串和錯誤
    """
    if not text or not text.strip():
        return text

    try:
        # Google Translate 有字數限制，通常是5000字元
        if len(text) > 4500:
            # 分段翻譯
            chunks = [text[i:i+4500] for i in range(0, len(text), 4500)]
            translated_chunks = []
            for chunk in chunks:
                translated_chunks.append(translator.translate(chunk))
                time.sleep(0.5)  # 避免請求過於頻繁
            return ''.join(translated_chunks)
        else:
            return translator.translate(text)
    except Exception as e:
        print(f"翻譯錯誤: {e}")
        return text  # 翻譯失敗則返回原文

def translate_paragraph(paragraph, translator):
    """
    翻譯段落內容，保留格式
    """
    if not paragraph.text.strip():
        return

    # 處理段落中的 runs（保留格式）
    full_text = paragraph.text
    translated_text = translate_text(full_text, translator)

    if not translated_text:
        return

    # 清除原有內容但保留格式
    # 簡單方法：保留第一個 run 的格式，清除其他
    if paragraph.runs:
        # 保存第一個 run 的格式
        first_run = paragraph.runs[0]

        # 清除所有 runs
        for run in paragraph.runs:
            run.text = ''

        # 在第一個 run 中設置翻譯後的文本
        first_run.text = translated_text
    else:
        paragraph.text = translated_text

def translate_table(table, translator):
    """
    翻譯表格內容，保留格式
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                translate_paragraph(paragraph, translator)

def translate_textboxes(element, translator):
    """
    翻譯元素中的所有文字框
    """
    # 查找所有文字框
    textboxes = element.findall('.//' + qn('w:txbxContent'))

    if not textboxes:
        return

    for textbox in textboxes:
        # 在文字框中查找所有文字元素
        text_elements = textbox.findall('.//' + qn('w:t'))

        for text_elem in text_elements:
            if text_elem.text and text_elem.text.strip():
                original_text = text_elem.text
                translated_text = translate_text(original_text, translator)
                if translated_text:
                    text_elem.text = translated_text
                    print(f"  翻譯文字框: {original_text[:30]}... -> {translated_text[:30]}...")
                time.sleep(0.2)  # 避免請求過於頻繁

def translate_document(input_file, output_file):
    """
    翻譯整個 Word 文件
    """
    print(f"正在載入文件: {input_file}")
    doc = Document(input_file)

    # 初始化翻譯器 (中文 -> 英文)
    translator = GoogleTranslator(source='zh-CN', target='en')

    print("開始翻譯段落...")
    # 翻譯所有段落
    total_paragraphs = len(doc.paragraphs)
    for i, paragraph in enumerate(doc.paragraphs, 1):
        if paragraph.text.strip():
            print(f"翻譯段落 {i}/{total_paragraphs}: {paragraph.text[:50]}...")
            translate_paragraph(paragraph, translator)
            time.sleep(0.3)  # 避免請求過於頻繁

    print("開始翻譯表格...")
    # 翻譯所有表格
    total_tables = len(doc.tables)
    for i, table in enumerate(doc.tables, 1):
        print(f"翻譯表格 {i}/{total_tables}...")
        translate_table(table, translator)
        time.sleep(0.3)

    print("開始翻譯文字框（流程圖）...")
    # 翻譯所有文字框
    for element in doc.element.body:
        if element.tag.endswith('p'):
            translate_textboxes(element, translator)

    # 儲存翻譯後的文件
    print(f"儲存翻譯後的文件: {output_file}")
    doc.save(output_file)
    print("翻譯完成！")

if __name__ == "__main__":
    input_file = "document.docx"
    output_file = "document_english.docx"

    translate_document(input_file, output_file)
