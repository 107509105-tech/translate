# translate_with_terms.py
# 功能：將 Word 文檔中的繁體中文翻譯成英文（使用 PCB 術語對照表）
# 特色：保留原始格式、處理表格/頁首頁尾/流程圖、智能合併段落、專業術語翻譯

import os
os.system('cls')  # 清空命令行畫面

# Word 文檔處理相關庫
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# LLM API 客戶端
from openai import OpenAI

# 其他工具庫
import re
import config
import time
from pathlib import Path
import json


# ==================== 載入術語對照表 ====================
def load_pcb_terms(json_path="pcb_terms_from_pdf.json"):
    """
    載入 PCB 術語對照表

    Args:
        json_path: JSON 檔案路徑

    Returns:
        dict: 術語對照字典
    """
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            terms = json.load(f)
        print(f"✓ 成功載入 {len(terms)} 個 PCB 術語")
        return terms
    except FileNotFoundError:
        print(f"⚠ 警告：找不到術語對照表文件 {json_path}，將不使用術語對照")
        return {}
    except json.JSONDecodeError as e:
        print(f"⚠ 警告：術語對照表格式錯誤 - {e}")
        return {}


def format_terms_for_prompt(terms_dict, max_terms=100):
    """
    將術語對照表格式化為適合放入 Prompt 的文字

    Args:
        terms_dict: 術語對照字典
        max_terms: 最多使用的術語數量（避免 Prompt 過長）

    Returns:
        str: 格式化後的術語對照文字
    """
    if not terms_dict:
        return ""

    # 選取前 max_terms 個術語
    selected_terms = list(terms_dict.items())[:max_terms]

    # 格式化為文字
    terms_text = "\n專業術語對照表（繁體中文 → 英文）：\n"
    for i, (english, translations) in enumerate(selected_terms, 1):
        traditional = translations.get("traditional", "")
        if traditional and english:
            terms_text += f"  - {traditional} → {english}\n"

    return terms_text


# 載入術語對照表
PCB_TERMS = load_pcb_terms()
TERMS_REFERENCE = format_terms_for_prompt(PCB_TERMS, max_terms=150)


# ==================== 設定區 ====================
# LLM 設定（使用 OpenAI 相容的 API）
client = OpenAI(
    api_key=config.LLM_API_KEY,       # API 金鑰
    base_url=config.LLM_API_BASE       # API 基礎 URL
)
LLM_MODEL_NAME = config.LLM_MODEL_NAME  # 模型名稱


# =====================  Prompt =====================
# LLM 翻譯系統提示詞（加入術語對照表）
TRANSLATE_SYSTEM_PROMPT = f""" 你是一位美國史丹佛大學語文學的博士及教授，精通半導體及印刷電路板等行業的英語及繁體中文的術語與專業表達，翻譯的著作及論文超過1000篇，
並曾主導IPC以及SEMICON West、SEMICON Taiwan等多個國際標準機構及協會，將其著作及規範由繁體中文翻譯成英語的工作。
我希望你能幫我將以下繁體中文翻譯成英語，並遵循以下規則：
(1) 翻譯時要準確傳達標準及規定以及方法。
(2) 翻譯時要保持專業術語的一致性，並確保符合國際標準。
(3) 翻譯時要注意語法和拼寫的正確性，確保符合英語語言習慣。
(4) 若翻譯為整句時，要符合語意，且符合句首大寫；若為片語、專有名詞、簡短標題，則只需符合字首大寫。
(5) 翻譯專業術語時，優先參考下方的「專業術語對照表」。
(6) 輸出翻譯字串，不要有多餘的說明或解釋，保持簡短。
(7) 符號保持一致，例如：【】。
(8) 所有形如「圖[中文數字]」的詞彙，必須翻譯為「Figure <對應的阿拉伯數字>」。

{TERMS_REFERENCE}

範例:
翻譯字串:中華精測科技股份有限公司
輸出:Chunghwa Precision Test Tech. Co., Ltd.

翻譯字串:耐磨性測試
輸出:Abrasion Resistance Test
"""
# ================================================


def starts_with_number(text):
    """
    檢查段落是否以數字開頭（章節標記）

    Args:
        text: 要檢查的文字

    Returns:
        bool: 如果以數字+點號開頭則返回 True（例如：1.、2.1）

    Note:
        使用正向預查 (?=\.) 確保數字後面跟著點號
    """
    text = text.lstrip()
    # 匹配開頭是數字的模式，例如：2.1, 1.2.3, 10, 等
    # (?=\.) 是正向預查，確保數字後面有點號
    return bool(re.match(r'^\d+(?=\.)', text))


def starts_with_multiple_spaces(text, min_spaces=2):
    """
    檢查段落開頭是否有多個空格

    Args:
        text: 要檢查的文字
        min_spaces: 最少空格數量，默認為 2

    Returns:
        bool: 如果開頭空格數 >= min_spaces 則返回 True

    Note:
        用於識別需要合併到章節的後續段落
    """
    match = re.match(r'^( +)', text)  # ( +) 捕獲一個或多個空格
    if match:
        # 檢查捕獲的空格數量是否達到最小要求
        return len(match.group(1)) >= min_spaces
    return False



def find_merge_blocks_by_spacing(paragraphs):
    """
    根據空格規則找出需要合併的段落區塊

    合併規則：
    1. 以數字開頭的段落（如 1.、2.1）作為章節開始
    2. 後續段落如果開頭有 2 個以上空格，則合併到該章節
    3. 直到遇到下一個章節開始，或空格數不足的段落為止

    Args:
        paragraphs: Word 文檔的段落列表

    Returns:
        list: 合併區塊的列表，每個元素為 (start_idx, end_idx) 元組

    Example:
        段落 0: "1. 標題"          ← 章節開始
        段落 1: "  這是內容"       ← 有 2 個空格，合併
        段落 2: "  繼續內容"       ← 有 2 個空格，合併
        段落 3: "不合併"           ← 無空格，停止合併
        → 返回 [(0, 2)]
    """
    merge_blocks = []  # 存儲合併區塊的起始和結束索引
    i = 0

    while i < len(paragraphs):
        para_text = paragraphs[i].text

        # 檢查是否為章節開始（以數字開頭）
        if starts_with_number(para_text):
            start_idx = i
            print(f"\n[{i}] 找到章節開始: {para_text[:10]}...")

            # 往後尋找需要合併的段落（有 2 個以上空格）
            j = i + 1
            while j < len(paragraphs):
                next_para_text = paragraphs[j].text

                # 遇到下一個章節，停止合併
                if starts_with_number(next_para_text):
                    break

                # 檢查是否有足夠的前導空格
                if starts_with_multiple_spaces(next_para_text, min_spaces=2):
                    print(f"[{j}] 合併此段落（有空格）: {next_para_text[:10]}...")
                    j += 1
                else:
                    # 空格不足，停止合併
                    break

            end_idx = j - 1

            # 只有當至少有 2 個段落時才記錄為合併區塊
            if end_idx > start_idx:
                merge_blocks.append((start_idx, end_idx))
                print(f"→ 合併區塊: 段落 {start_idx} 到 {end_idx} ({end_idx - start_idx + 1} 個段落)")

            i = j  # 跳到下一個未處理的段落
        else:
            i += 1

    return merge_blocks


def is_chinese(text):
    """
    檢查文字中是否包含中文字符

    Args:
        text: 要檢查的文字

    Returns:
        bool: 如果包含中文字符則返回 True
    """
    return bool(re.search('[\u4e00-\u9fff]', text))  # Unicode 中文範圍


def translate_to_english(text):
    """
    使用 LLM 將繁體中文翻譯成英文

    Args:
        text: 要翻譯的文字

    Returns:
        str: 翻譯後的英文文字

    Note:
        - 保留原始的前導空格
        - 空文字或無中文的文字直接返回
        - 使用 temperature=0.0 確保翻譯的一致性
    """
    # 空文字或無中文的直接跳過
    if not text or not text.strip() or not is_chinese(text):
        return text

    # 記錄原始前導空格數量，翻譯後需要保留
    leading_spaces = len(text) - len(text.lstrip())

    # 呼叫 LLM API 進行翻譯
    response = client.chat.completions.create(
        model=LLM_MODEL_NAME,
        temperature=0.0,  # 設為 0 以獲得確定性的翻譯結果
        messages=[
            {"role": "system", "content": TRANSLATE_SYSTEM_PROMPT},
            {"role": "user",   "content": f"翻譯字串: {text}\n輸出:"}
        ]
    )

    # 提取翻譯結果並恢復前導空格
    translated = response.choices[0].message.content.strip()
    translated = ' ' * leading_spaces + translated

    return translated


def clear_run_text_preserve_drawing(run):
    """
    清空 run 內所有文字節點，但保留圖片/繪圖等其他子節點

    Args:
        run: Word 文檔的 Run 對象

    Note:
        這樣可以避免刪除段落中的圖片、繪圖等非文字元素
        只清空 <w:t> 文字節點的內容
    """
    r = run._r  # 取得底層 XML 元素
    for t in r.findall(qn('w:t')):  # 找到所有文字節點
        t.text = ''  # 清空文字內容


def set_run_text_preserve_drawing(run, text):
    """
    在現有 run 的文字節點寫入文字，若沒有文字節點則新增一個

    Args:
        run: Word 文檔的 Run 對象
        text: 要寫入的文字

    Note:
        - 優先使用第一個文字節點
        - 清空其他多餘的文字節點
        - 保留圖片、繪圖等非文字元素
    """
    r = run._r  # 取得底層 XML 元素
    t_elems = r.findall(qn('w:t'))  # 找到所有文字節點

    if t_elems:
        # 如果有文字節點，使用第一個
        t_elems[0].text = text
        # 清空其他多餘的文字節點
        for extra in t_elems[1:]:
            extra.text = ''
    else:
        # 如果沒有文字節點，新增一個
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)



# ===================== [5] 判斷是否為特殊格式 run =====================
def is_special_format(run):
    """
    判斷 run 是否具有特殊格式（粗體、斜體、底線、顏色等）

    Args:
        run: Word 文檔的 Run 對象

    Returns:
        bool: 如果有特殊格式則返回 True

    Note:
        此函數目前未被使用，保留作為未來擴展用途
    """
    font = run.font
    # 檢查粗體、斜體、底線
    if font.bold or font.italic or font.underline:
        return True
    # 檢查字體顏色
    if font.color and font.color.rgb and str(font.color.rgb) not in ('000000', None, ''):
        return True
    # 檢查螢光筆顏色
    if font.highlight_color:
        return True
    return False


def translate_paragraph(paragraph):
    """
    翻譯單個段落的文字

    Args:
        paragraph: Word 文檔的段落對象

    Note:
        - 跳過空白段落
        - 保留段落中的圖片和繪圖元素
        - 將翻譯結果寫入第一個 run，清空其他 run 的文字
    """
    # 跳過空白段落
    if not paragraph.text.strip():
        return

    # 取得完整段落文字並翻譯
    full_text = paragraph.text
    print(full_text)
    translated_text = translate_to_english(full_text)

    if not translated_text:
        return

    # 清空所有 Run 的文字節點，再把翻譯後的文字寫回第一個 Run
    # 這樣可以保留圖片與繪圖等非文字元素
    if paragraph.runs:
        first_run = paragraph.runs[0]
        for run in paragraph.runs:
            clear_run_text_preserve_drawing(run)
        set_run_text_preserve_drawing(first_run, translated_text)
    else:
        # 如果沒有 run，直接設定段落文字
        paragraph.text = translated_text


def translate_table(table):
    """
    翻譯表格中所有儲存格的文字

    Args:
        table: Word 文檔的表格對象

    Note:
        遍歷表格的每一行、每一個儲存格、每一個段落進行翻譯
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                translate_paragraph(paragraph)


def translate_header_footer_full(doc):
    """
    翻譯頁首頁尾的所有內容（包括文字和表格）

    Args:
        doc: Word 文檔對象

    Note:
        - 處理所有類型的頁首頁尾：一般、首頁、偶數頁
        - 避免重複翻譯已連結到前一區段的頁首頁尾
        - 翻譯段落文字和表格內容
    """
    for section in doc.sections:

        # 收集所有可能的頁首
        headers = [
            section.header,              # 一般頁首
            section.first_page_header,   # 首頁頁首
            section.even_page_header     # 偶數頁頁首
        ]
        # 收集所有可能的頁尾
        footers = [
            section.footer,              # 一般頁尾
            section.first_page_footer,   # 首頁頁尾
            section.even_page_footer     # 偶數頁頁尾
        ]

        # 一起處理所有頁首頁尾
        for hf in headers + footers:
            # 跳過不存在或連結到前一區段的頁首頁尾
            # （連結到前一區段表示內容相同，避免重複翻譯）
            if not hf or hf.is_linked_to_previous:
                continue

            # 1. 翻譯頁首頁尾內的一般段落文字
            for para in hf.paragraphs:
                if para.text.strip() and is_chinese(para.text):
                    translated = translate_to_english(para.text)
                    if translated and para.runs:
                        # 清空所有 run 的文字
                        for r in para.runs:
                            r.text = ''
                        # 將翻譯結果寫入第一個 run
                        para.runs[0].text = translated

            # 2. 翻譯頁首頁尾內的表格內容
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip() and is_chinese(para.text):
                                translated = translate_to_english(para.text)
                                if translated and para.runs:
                                    # 清空所有 run 的文字
                                    for r in para.runs:
                                        r.text = ''
                                    # 將翻譯結果寫入第一個 run
                                    para.runs[0].text = translated


# ===================== 翻譯流程圖文字函式 =====================
def translate_textboxes_in_doc(doc):
    """
    翻譯文檔中所有文字方塊（流程圖、圖表等）的文字

    Args:
        doc: Word 文檔對象

    Note:
        - 翻譯文字方塊內的所有文字
        - 自動調整翻譯後的字體大小、行高和對齊方式
        - 確保翻譯後的英文文字能正確顯示在流程圖中
    """
    body = doc.element.body
    if body is None:
        return

    # 找到所有文字方塊內容（w:txbxContent）
    textboxes = body.findall('.//' + qn('w:txbxContent'))

    if not textboxes:
        return

    for textbox in textboxes:
        # 先翻譯所有文字元素
        text_elements = textbox.findall('.//' + qn('w:t'))
        for text_elem in text_elements:
            if text_elem.text and text_elem.text.strip():
                original_text = text_elem.text
                # 翻譯文字
                translated_text = translate_to_english(original_text)
                if translated_text:
                    text_elem.text = translated_text

                # 如果原始文字包含中文，需要調整格式
                if is_chinese(original_text):
                    half_points = 11  # 字體大小（半磅，11 = 5.5pt）

                    # 調整字體大小
                    for r in textbox.findall('.//' + qn('w:r')):
                        rPr = r.find(qn('w:rPr'))
                        if rPr is None:
                            rPr = OxmlElement('w:rPr')
                            r.append(rPr)
                        sz = rPr.find(qn('w:sz'))
                        if sz is None:
                            sz = OxmlElement('w:sz')
                            rPr.append(sz)
                        sz.set(qn('w:val'), str(half_points))

                    # 調整行高
                    for p in textbox.findall('.//' + qn('w:p')):
                        pPr = p.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            p.insert(0, pPr)
                        spacing = pPr.find(qn('w:spacing'))
                        if spacing is None:
                            spacing = OxmlElement('w:spacing')
                            pPr.append(spacing)
                        spacing.set(qn('w:line'), '130')  # 行高 130 twips
                        spacing.set(qn('w:lineRule'), 'exact')  # 固定行高

                    # 調整文字對齊方式為置中
                    for p in textbox.findall('.//' + qn('w:p')):
                        pPr = p.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            p.insert(0, pPr)

                        # 設定為置中對齊
                        jc = pPr.find(qn('w:jc'))
                        if jc is None:
                            jc = OxmlElement('w:jc')
                            pPr.append(jc)
                        jc.set(qn('w:val'), 'center')


# ===================== 縮小表格內英文文字函式 =====================
def shrink_table_english_font(table, ratio=0.82):
    """
    縮小表格內純英文文字的字體大小

    Args:
        table: Word 文檔的表格對象
        ratio: 縮小比例，默認為 0.82 (縮小至 82%)

    Note:
        - 只處理純英文段落，跳過包含中文的段落
        - 避免中英混排時誤判
        - 同步調整 w:sz、w:szCs、w:szFarEast 確保字體大小一致
        - 最小字體大小為 10pt (20 半磅)
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                text = paragraph.text.strip()
                # 關鍵：只要段落包含中文就跳過（避免中英混排誤判）
                if not text or any('\u4e00' <= c <= '\u9fff' for c in text):
                    continue

                for run in paragraph.runs:
                    if not run.text.strip():
                        continue

                    r = run._r  # 底層 XML 元素 <w:r>
                    rPr = r.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        r.insert(0, rPr)

                    # === 處理 w:sz (西文字體大小) ===
                    sz = rPr.find(qn('w:sz'))
                    if sz is None:
                        sz = OxmlElement('w:sz')
                        rPr.insert(0, sz)

                    # 計算新的字體大小
                    if sz.get(qn('w:val')):
                        current = int(sz.get(qn('w:val')))
                        new_val = max(20, int(current * ratio))  # 最低 10pt (20 半磅)
                    else:
                        new_val = max(20, int(22 * ratio))  # 預設 11pt 縮小至 9pt

                    sz.set(qn('w:val'), str(new_val))

                    # === 處理 w:szCs (複雜字體大小，必須同步！) ===
                    # 中日韓等複雜字體大小設定
                    szCs = rPr.find(qn('w:szCs'))
                    if szCs is None:
                        szCs = OxmlElement('w:szCs')
                        rPr.insert(0, szCs)
                    szCs.set(qn('w:val'), str(new_val))

                    # === 處理 w:szFarEast (東亞字體大小) ===
                    # 舊版 Word 使用的東亞字體設定
                    szFarEast = rPr.find(qn('w:szFarEast'))
                    if szFarEast is None:
                        szFarEast = OxmlElement('w:szFarEast')
                        rPr.insert(0, szFarEast)
                    szFarEast.set(qn('w:val'), str(new_val))


# ===================== 強制 Times New Roman 字體函式 =====================
def force_times_new_roman(doc):
    """
    強制全文件所有文字改成 Times New Roman 字體

    Args:
        doc: Word 文檔對象

    Note:
        - 處理正文、表格、流程圖、頁首頁尾的所有文字
        - ASCII 和西文字體設為 Times New Roman
        - 東亞字體設為細明體（MingLiU）
        - 確保文檔的統一字體風格
    """
    def _set_font(run):
        """
        設定單個 run 的字體

        Args:
            run: Word 文檔的 Run 對象
        """
        r = run._r
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)

        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)

        # 設定各種字體類型
        rFonts.set(qn('w:ascii'), 'Times New Roman')    # ASCII 字符字體
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')    # 高位 ANSI 字符字體
        rFonts.set(qn('w:cs'), 'Times New Roman')       # 複雜字體（Complex Script）
        rFonts.set(qn('w:eastAsia'), '細明體')          # 東亞字體（中文）

    # ===== 處理正文段落 =====
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _set_font(run)

    # ===== 處理正文表格 =====
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_font(run)

    # ===== 處理頁首頁尾 =====
    for section in doc.sections:
        # 檢查所有類型的頁首頁尾
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            if hf:
                # 處理頁首頁尾的段落
                for paragraph in hf.paragraphs:
                    for run in paragraph.runs:
                        _set_font(run)
                # 處理頁首頁尾的表格
                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    _set_font(run)

    # ===== 處理流程圖文字框 =====
    for textbox in doc.element.body.findall('.//' + qn('w:txbxContent')):
        for run in textbox.findall('.//' + qn('w:r')):
            rPr = run.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                run.insert(0, rPr)

            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)

            # 設定流程圖文字框的字體
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:cs'), 'Times New Roman')
            rFonts.set(qn('w:eastAsia'), '細明體')


# ==================== 主翻譯函式 ====================
def translate_document(input_file, output_file):
    """
    翻譯 Word 文檔的主函數（使用 PCB 術語對照表）

    Args:
        input_file: 輸入的 Word 文檔路徑
        output_file: 輸出的翻譯後 Word 文檔路徑

    翻譯流程：
        1. 載入文檔
        2. 找出需要合併的段落區塊（章節開頭 + 後續有空格的段落）
        3. 翻譯合併後的段落（使用術語對照表）
        4. 翻譯一般段落
        5. 翻譯表格內容
        6. 翻譯頁首頁尾
        7. 翻譯流程圖文字框
        8. 調整表格英文字體大小
        9. 統一設定字體為 Times New Roman
        10. 儲存翻譯結果

    Note:
        - 保留原始格式和圖片
        - 智能合併章節段落以提高翻譯品質
        - 使用 PCB 專業術語對照表確保術語翻譯準確
        - 自動調整字體大小和樣式
    """
    print(f"="*90)
    print(f"載入檔案：{input_file}\n")
    doc = Document(input_file)

    # ===== 步驟 1: 找出需要合併的段落區塊 =====
    merge_blocks = find_merge_blocks_by_spacing(doc.paragraphs)

    # 建立合併區塊的字典
    # key: 起始段落索引, value: (合併後的文字, 結束段落索引)
    merge_dict = {}
    for start_idx, end_idx in merge_blocks:
        merged_text = ""
        for j in range(start_idx, end_idx + 1):
            if j < len(doc.paragraphs):
                merged_text += doc.paragraphs[j].text
        merge_dict[start_idx] = (merged_text, end_idx)

    # 建立需要清空的段落索引集合
    # （合併區塊內的非起始段落需要清空，因為內容已合併到起始段落）
    paragraphs_to_clear = set()
    for start_idx, end_idx in merge_blocks:
        for j in range(start_idx + 1, end_idx + 1):
            paragraphs_to_clear.add(j)

    # ===== 步驟 2-4: 翻譯段落 =====
    print(f"\n開始翻譯段落（使用 PCB 術語對照表）...")
    for i, paragraph in enumerate(doc.paragraphs):
        # 情況 1: 需要清空的段落（合併區塊內的後續段落）
        if i in paragraphs_to_clear:
            # 清空文字但保留格式和圖表
            # 因為這些段落的內容已經合併到起始段落翻譯了
            if paragraph.runs:
                for run in paragraph.runs:
                    clear_run_text_preserve_drawing(run)
            else:
                paragraph.text = ""
            continue

        # 情況 2: 合併區塊的起始段落
        if i in merge_dict:
            merged_text, end_idx = merge_dict[i]
            print(f"[{i}] 翻譯合併區塊 ({i}~{end_idx}): {merged_text[:50]}...")

            # 翻譯合併後的完整文字
            if merged_text.strip() and is_chinese(merged_text):
                translated_text = translate_to_english(merged_text)
                print(f"翻譯結果: {translated_text[:50]}...")

                if translated_text:
                    # 將翻譯結果寫入起始段落
                    if paragraph.runs:
                        first_run = paragraph.runs[0]
                        for run in paragraph.runs:
                            clear_run_text_preserve_drawing(run)
                        set_run_text_preserve_drawing(first_run, translated_text)
                    else:
                        paragraph.text = translated_text
            continue

        # 情況 3: 一般段落直接翻譯
        if paragraph.text.strip():
            print(f"[{i}] 翻譯一般段落: {paragraph.text[:50]}...")
            translate_paragraph(paragraph)


    # ===== 步驟 5: 翻譯表格內容 =====
    print("\n開始翻譯表格內容...")
    for t_idx, table in enumerate(doc.tables, 1):
        translate_table(table)

    # ===== 步驟 6: 翻譯頁首頁尾 =====
    print("\n開始翻譯頁首頁尾...")
    translate_header_footer_full(doc)

    # ===== 步驟 7: 翻譯流程圖文字框 =====
    print(f"翻譯流程圖.....")
    translate_textboxes_in_doc(doc)

    # ===== 步驟 8: 調整表格英文字體大小 =====
    print("開始縮小表格內英文字體（82%）...")
    # 處理正文表格
    for table in doc.tables:
        shrink_table_english_font(table, ratio=0.82)

    # 處理頁首頁尾內的表格
    for section in doc.sections:
        for hf in (section.header, section.footer,
                section.first_page_header, section.first_page_footer,
                section.even_page_header, section.even_page_footer):
            if hf and hf.tables:
                for t in hf.tables:
                    shrink_table_english_font(t, ratio=0.82)

    # ===== 步驟 9: 統一設定字體為 Times New Roman =====
    print("強制全文件字體為 Times New Roman...")
    force_times_new_roman(doc)

    # ===== 步驟 10: 儲存翻譯結果 =====
    print(f"="*90)
    print(f"儲存翻譯結果 → {output_file}")
    doc.save(output_file)
    print("翻譯完成！")


# ==================== 一鍵執行 ====================
if __name__ == "__main__":

    print(f'多語文檔轉譯專案（使用 PCB 術語對照表）\n')

    # 導入 Windows COM 客戶端（用於處理舊版 .doc 格式）
    import win32com.client

    def convert_doc_to_docx(doc_path):
        """
        將舊版 .doc 格式轉換為 .docx 格式

        Args:
            doc_path: .doc 文件的路徑

        Returns:
            Path: 轉換後的 .docx 文件路徑

        Note:
            - 使用 Word COM 客戶端進行轉換
            - 需要安裝 Microsoft Word
            - 轉換後會自動關閉 Word 應用程式
        """
        doc_path = Path(doc_path)
        docx_path = doc_path.with_suffix(".docx")

        # 啟動 Word 應用程式（隱藏視窗）
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        # 開啟 .doc 文件並另存為 .docx
        doc = word.Documents.Open(str(doc_path.absolute()))
        doc.SaveAs(str(docx_path.absolute()), FileFormat=16)  # 16 = .docx 格式
        doc.Close()
        word.Quit()

        return docx_path

    # ===== 可選功能：關閉背景 Word 程序 =====
    # import kill_all_word
    # kill_all_word.kill_word()   # 用 win32com.client 前建議先執行這一行關閉 Word 背景程序

    # ===== 可選功能：將 .doc 轉換為 .docx =====
    # f = Path((r"I:/translate_doc/W-QA-A070 ver02 原子吸收光譜儀(280FS AA)操作指導書.docx"))
    # f = convert_doc_to_docx(f)

    # ===== 可選功能：製作測試副本 =====
    # import make_test_copy
    # test_file = make_test_copy.make_fresh_copy(f, f.with_stem(f.stem + "_test"))
    # test_file = make_test_copy.make_fresh_copy()

    # ===== 設定輸入和輸出文件 =====
    # 輸入文件：要翻譯的 Word 文檔
    input_file = r"Q:\14_人工智慧處\01_AI部\01_智慧代理應用課\01_個人資料夾\Charles\W-QA-B002_OQC檢驗作業標準 (20251125提供QA內部軟體測試使用).docx"

    # 輸出文件：翻譯後的 Word 文檔
    output_file = "merged_translated_with_terms.docx"

    # output_file = test_file.with_stem(test_file.stem.replace("test", "output_[05]"))

    # ===== 執行翻譯並計時 =====
    start_time = time.time()

    translate_document(input_file, output_file)

    total_time = time.time() - start_time
    print(f"總耗時：{total_time:.2f} 秒（{total_time/60:.2f} 分鐘）")

    # ===== 記錄翻譯時間到日誌文件 =====
    from datetime import datetime
    with open("翻譯時間紀錄_with_terms.log", "a", encoding="utf-8") as f:
        f.write(f"{datetime.now():%Y-%m-%d %H:%M:%S} | [with_terms] | 總耗時: {total_time/60:.2f} 分鐘\n")

    # ===== 可選功能：刪除測試檔案 =====
    # test_file.unlink()  # 清理測試檔案
