#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
檢查 Word 文件的結構
"""

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

def check_document_structure(input_file):
    """
    檢查文件結構，包括段落、表格、文字框等
    """
    print(f"正在檢查文件: {input_file}\n")
    doc = Document(input_file)

    # 檢查段落
    print(f"段落數量: {len(doc.paragraphs)}")
    for i, para in enumerate(doc.paragraphs[:10], 1):
        if para.text.strip():
            print(f"  段落 {i}: {para.text[:50]}...")

    # 檢查表格
    print(f"\n表格數量: {len(doc.tables)}")

    # 檢查文字框和圖形
    print("\n檢查文字框和圖形...")

    # 遍歷文件中的所有元素
    for element in doc.element.body:
        # 檢查段落中的文字框
        if element.tag.endswith('p'):
            # 查找段落中的所有 drawing 元素
            drawings = element.findall('.//' + qn('w:drawing'))
            if drawings:
                print(f"找到 {len(drawings)} 個繪圖元素")

            # 查找文字框
            textboxes = element.findall('.//' + qn('w:txbxContent'))
            if textboxes:
                print(f"找到 {len(textboxes)} 個文字框")
                for tb in textboxes:
                    # 提取文字框中的文字
                    text_elements = tb.findall('.//' + qn('w:t'))
                    for te in text_elements:
                        if te.text:
                            print(f"  文字框內容: {te.text[:50]}...")

if __name__ == "__main__":
    check_document_structure("document.docx")
