#!/usr/bin/env python3
from docx import Document
from deep_translator import GoogleTranslator
import time
import re

def translate_text(text, source='zh-TW', target='en'):
    """Translate text from Chinese to English"""
    if not text.strip():
        return text

    try:
        translator = GoogleTranslator(source=source, target=target)
        # Split long text into chunks if needed (Google Translate has a limit)
        max_length = 4900
        if len(text) <= max_length:
            return translator.translate(text)
        else:
            # Split by sentences and translate in chunks
            chunks = []
            current_chunk = ""
            for sentence in text.split('。'):
                if len(current_chunk) + len(sentence) + 1 <= max_length:
                    current_chunk += sentence + '。'
                else:
                    if current_chunk:
                        chunks.append(current_chunk)
                    current_chunk = sentence + '。'
            if current_chunk:
                chunks.append(current_chunk)

            translated_chunks = []
            for chunk in chunks:
                translated_chunks.append(translator.translate(chunk))
                time.sleep(0.5)  # Rate limiting

            return ' '.join(translated_chunks)
    except Exception as e:
        print(f"Translation error: {e}")
        return text

def starts_with_number(text):
    """檢查段落是否以數字開頭（章節標記）"""
    text = text.lstrip()
    # 匹配開頭是數字的模式，例如：4.9, 1.2.3, 10, 等
    return bool(re.match(r'^\d+', text))

def starts_with_multiple_spaces(text, min_spaces=2):
    """檢查段落開頭是否有多個空格（至少 min_spaces 個）"""
    match = re.match(r'^( +)', text)
    if match:
        return len(match.group(1)) >= min_spaces
    return False

def find_merge_blocks_by_spacing(paragraphs):
    """
    根據新規則找出需要合併的段落區塊：
    - 如果段落以數字開頭（章節），則作為新區塊的開始
    - 後續段落如果開頭有 2 個以上空格，則合併到該區塊
    - 直到遇到下一個以數字開頭的段落，或開頭空格少於 2 個的段落

    Returns:
        List of tuples (start_index, end_index) for paragraphs to merge
    """
    merge_blocks = []
    i = 0

    while i < len(paragraphs):
        para_text = paragraphs[i].text

        # 檢查是否以數字開頭（章節開始）
        if starts_with_number(para_text):
            start_idx = i
            print(f"\n[{i}] 找到章節開始: {para_text[:60]}...")

            # 繼續往下找，合併所有開頭有 2+ 空格的段落
            j = i + 1
            while j < len(paragraphs):
                next_para_text = paragraphs[j].text

                # 如果遇到下一個章節（以數字開頭），停止
                if starts_with_number(next_para_text):
                    break

                # 如果開頭有 2+ 空格，繼續合併
                if starts_with_multiple_spaces(next_para_text, min_spaces=2):
                    print(f"[{j}] 合併此段落（有空格）: {next_para_text[:60]}...")
                    j += 1
                else:
                    # 開頭沒有足夠空格，停止合併
                    break

            end_idx = j - 1

            # 如果有需要合併的段落（不只一個段落）
            if end_idx > start_idx:
                merge_blocks.append((start_idx, end_idx))
                print(f"→ 合併區塊: 段落 {start_idx} 到 {end_idx} ({end_idx - start_idx + 1} 個段落)")

            i = j  # 跳到下一個未處理的段落
        else:
            i += 1

    return merge_blocks

# Read the original document
doc = Document('worksheet.docx')

print("=" * 60)
print("分析文檔結構...")
print("=" * 60)

# 先顯示所有段落，方便理解
print("\n原始段落預覽：")
for i, para in enumerate(doc.paragraphs[:30]):  # 只顯示前 30 個段落
    text = para.text
    if text.strip():
        # 顯示空格
        display_text = text.replace(' ', '·')  # 用 · 標記空格
        print(f"[{i:2d}] {display_text[:80]}")

print("\n" + "=" * 60)
print("開始尋找合併區塊...")
print("=" * 60)

merge_blocks = find_merge_blocks_by_spacing(doc.paragraphs)

print("\n" + "=" * 60)
print(f"共找到 {len(merge_blocks)} 個需要合併的區塊")
print("=" * 60)

# Create a new document for the translation
new_doc = Document()

# Process paragraphs with merge logic
paragraphs_to_process = []
i = 0
while i < len(doc.paragraphs):
    # Check if current index is the start of a merge block
    merge_found = False
    for start_idx, end_idx in merge_blocks:
        if i == start_idx:
            # Merge paragraphs from start_idx to end_idx
            merged_text = ""
            for j in range(start_idx, end_idx + 1):
                if j < len(doc.paragraphs):
                    merged_text += doc.paragraphs[j].text
            paragraphs_to_process.append((merged_text, doc.paragraphs[i]))
            i = end_idx + 1  # Skip to after the merged section
            merge_found = True
            break

    if not merge_found:
        paragraphs_to_process.append((doc.paragraphs[i].text, doc.paragraphs[i]))
        i += 1

# Translate and add to new document
print("\n" + "=" * 60)
print("開始翻譯...")
print("=" * 60)

for idx, (text, original_para) in enumerate(paragraphs_to_process):
    print(f"\n翻譯段落 {idx + 1}/{len(paragraphs_to_process)}...")

    if text.strip():
        translated_text = translate_text(text)
        # Add paragraph to new document
        new_para = new_doc.add_paragraph(translated_text)

        # Copy formatting from original paragraph
        new_para.style = original_para.style

        print(f"  原文: {text[:50]}...")
        print(f"  譯文: {translated_text[:50]}...")
    else:
        # Keep empty paragraphs
        new_doc.add_paragraph("")

    time.sleep(0.5)  # Rate limiting

# Save the translated document
output_filename = 'worksheet_english_v2.docx'
new_doc.save(output_filename)
print("\n" + "=" * 60)
print(f"翻譯完成！已儲存到 {output_filename}")
print("=" * 60)
